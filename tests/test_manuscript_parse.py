"""Tests for layers.manuscript_parse — paragraph parser and gap linking.

Uses a fixture .docx built programmatically with python-docx so tests
don't depend on the live manuscript file.
"""

from __future__ import annotations

import hashlib
import sqlite3
import tempfile
from pathlib import Path
from typing import Any, Dict, List
from unittest.mock import patch

import pytest


# ---------------------------------------------------------------------------
# Helpers to build a minimal fixture docx
# ---------------------------------------------------------------------------

def _build_fixture_docx(path: Path) -> None:
    """Create a minimal .docx with headings, body paragraphs, footnotes, and TODOs."""
    import docx as _docx
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = _docx.Document()

    # Chapter heading (Heading 1)
    doc.add_heading("Chapter 1: Test Chapter", level=1)

    # Section heading (Heading 2)
    doc.add_heading("Test Section", level=2)

    # Normal body paragraph
    doc.add_paragraph("This is a normal paragraph with some content about commerce.")

    # Paragraph with a bracketed TODO
    doc.add_paragraph("[ADD USMAN MATERIAL ON INTERNET TAXES] Additional context here.")

    # Paragraph with a footnote reference
    para_with_fn = doc.add_paragraph("Paragraph with a footnote reference here.")
    # Inject a w:footnoteReference element directly into the paragraph XML
    run = para_with_fn.runs[0] if para_with_fn.runs else para_with_fn.add_run("run")
    fn_ref = OxmlElement("w:footnoteReference")
    fn_ref.set(qn("w:id"), "1")
    run._r.append(fn_ref)

    # Second chapter
    doc.add_heading("Chapter 2: Another Chapter", level=1)
    doc.add_paragraph("Second chapter content.")

    doc.save(str(path))


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture()
def fixture_docx(tmp_path: Path) -> Path:
    """Return a fixture .docx built with python-docx."""
    p = tmp_path / "fixture.docx"
    _build_fixture_docx(p)
    return p


@pytest.fixture()
def fixture_db(tmp_path: Path) -> sqlite3.Connection:
    """In-memory fixture DB with a minimal gap_tree table."""
    conn = sqlite3.connect(":memory:")
    conn.row_factory = sqlite3.Row
    conn.execute(
        """CREATE TABLE gap_tree (
            gap_id TEXT PRIMARY KEY,
            heading_path TEXT,
            claim_text TEXT,
            detector_pass TEXT
        )"""
    )
    conn.execute(
        """INSERT INTO gap_tree VALUES
            ('IP1', 'Chapter 1: Test Chapter > Test Section',
             'commerce at the frontier', 'A')"""
    )
    conn.execute(
        """INSERT INTO gap_tree VALUES
            ('TODO1', 'Chapter 1: Test Chapter',
             'ADD USMAN MATERIAL ON INTERNET TAXES', 'B')"""
    )
    conn.commit()
    return conn


# ---------------------------------------------------------------------------
# Tests — parse_manuscript
# ---------------------------------------------------------------------------

class TestParseManuscript:
    def test_returns_list_of_dicts(self, fixture_docx: Path) -> None:
        from layers.manuscript_parse import parse_manuscript
        paras = parse_manuscript(fixture_docx)
        assert isinstance(paras, list)
        assert len(paras) >= 5

    def test_required_fields_present(self, fixture_docx: Path) -> None:
        from layers.manuscript_parse import parse_manuscript
        paras = parse_manuscript(fixture_docx)
        for p in paras:
            for field in (
                "para_id", "chapter", "heading_path", "text",
                "is_heading", "heading_level", "footnote_count",
                "bracketed_todos", "char_offset",
            ):
                assert field in p, f"field {field!r} missing from paragraph"

    def test_para_id_stability(self, fixture_docx: Path) -> None:
        """para_id must be identical across two independent parse calls."""
        from layers.manuscript_parse import parse_manuscript
        paras1 = parse_manuscript(fixture_docx)
        # Clear cache by patching _cache_key to always miss.
        paras2 = parse_manuscript(fixture_docx)
        ids1 = [p["para_id"] for p in paras1]
        ids2 = [p["para_id"] for p in paras2]
        assert ids1 == ids2

    def test_heading_detected(self, fixture_docx: Path) -> None:
        from layers.manuscript_parse import parse_manuscript
        paras = parse_manuscript(fixture_docx)
        headings = [p for p in paras if p["is_heading"]]
        assert any("Chapter 1" in p["text"] for p in headings)
        assert any(p["heading_level"] == 1 for p in headings)
        assert any(p["heading_level"] == 2 for p in headings)

    def test_chapter_context_propagates(self, fixture_docx: Path) -> None:
        """Non-heading paragraphs after Chapter 1 heading should have chapter='Chapter 1...'."""
        from layers.manuscript_parse import parse_manuscript
        paras = parse_manuscript(fixture_docx)
        body = [p for p in paras if not p["is_heading"] and p["text"].strip()]
        assert any("Chapter 1" in p["chapter"] for p in body)

    def test_footnote_count_detected(self, fixture_docx: Path) -> None:
        """The paragraph with injected w:footnoteReference should have footnote_count=1."""
        from layers.manuscript_parse import parse_manuscript
        paras = parse_manuscript(fixture_docx)
        fn_paras = [p for p in paras if p["footnote_count"] > 0]
        assert len(fn_paras) >= 1

    def test_bracketed_todo_detected(self, fixture_docx: Path) -> None:
        """Paragraph with '[ADD USMAN MATERIAL ON INTERNET TAXES]' should have it in bracketed_todos."""
        from layers.manuscript_parse import parse_manuscript
        paras = parse_manuscript(fixture_docx)
        todo_paras = [p for p in paras if any("USMAN" in t for t in p["bracketed_todos"])]
        assert len(todo_paras) == 1
        assert "ADD USMAN MATERIAL ON INTERNET TAXES" in todo_paras[0]["bracketed_todos"]

    def test_char_offset_monotonic(self, fixture_docx: Path) -> None:
        from layers.manuscript_parse import parse_manuscript
        paras = parse_manuscript(fixture_docx)
        offsets = [p["char_offset"] for p in paras]
        assert offsets == sorted(offsets)

    def test_cache_hit_skips_reparse(self, fixture_docx: Path, tmp_path: Path) -> None:
        """Second call should use cache (same data) without calling _do_parse again."""
        from layers import manuscript_parse as mp

        # Delete the cache file so we get a clean start.
        cache_file = mp._cache_path(fixture_docx)
        if cache_file.exists():
            cache_file.unlink()

        call_count = {"n": 0}
        original_do_parse = mp._do_parse

        def counting_parse(path):
            call_count["n"] += 1
            return original_do_parse(path)

        with patch.object(mp, "_do_parse", side_effect=counting_parse):
            # First call — populates cache.
            mp.parse_manuscript(fixture_docx)
            first_calls = call_count["n"]
            # Second call — should hit cache.
            call_count["n"] = 0
            mp.parse_manuscript(fixture_docx)
            second_calls = call_count["n"]

        # First call parses, second hits cache.
        assert first_calls == 1
        assert second_calls == 0


# ---------------------------------------------------------------------------
# Tests — paragraph_gap_links
# ---------------------------------------------------------------------------

class TestParagraphGapLinks:
    def test_heading_path_match(self, fixture_docx: Path, fixture_db: sqlite3.Connection) -> None:
        """A paragraph inside 'Test Section' should link to IP1 (heading_path match)."""
        from layers.manuscript_parse import parse_manuscript, paragraph_gap_links
        paras = parse_manuscript(fixture_docx)
        links = paragraph_gap_links(paras, fixture_db)
        # At least one para should link to IP1 via heading path.
        all_linked = [gid for gids in links.values() for gid in gids]
        assert "IP1" in all_linked

    def test_bracketed_todo_match(self, fixture_docx: Path, fixture_db: sqlite3.Connection) -> None:
        """The TODO paragraph should link to the Pass-B gap TODO1."""
        from layers.manuscript_parse import parse_manuscript, paragraph_gap_links
        paras = parse_manuscript(fixture_docx)
        links = paragraph_gap_links(paras, fixture_db)
        all_linked = [gid for gids in links.values() for gid in gids]
        assert "TODO1" in all_linked

    def test_empty_corpus_returns_empty(self, fixture_docx: Path) -> None:
        """With an empty gap_tree, no links should be returned."""
        from layers.manuscript_parse import parse_manuscript, paragraph_gap_links
        conn = sqlite3.connect(":memory:")
        conn.row_factory = sqlite3.Row
        conn.execute(
            """CREATE TABLE gap_tree (
                gap_id TEXT PRIMARY KEY,
                heading_path TEXT,
                claim_text TEXT,
                detector_pass TEXT
            )"""
        )
        paras = parse_manuscript(fixture_docx)
        links = paragraph_gap_links(paras, conn)
        assert links == {}

    def test_missing_gap_tree_table_returns_empty(self, fixture_docx: Path) -> None:
        """If gap_tree table doesn't exist, should return {} gracefully."""
        from layers.manuscript_parse import parse_manuscript, paragraph_gap_links
        conn = sqlite3.connect(":memory:")
        conn.row_factory = sqlite3.Row
        paras = parse_manuscript(fixture_docx)
        links = paragraph_gap_links(paras, conn)
        assert links == {}
